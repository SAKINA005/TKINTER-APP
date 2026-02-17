"""
TEMPLATES D'EXPORT PROFESSIONNELS - 
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from datetime import datetime
from typing import Dict


def export_to_word_masterclass(data_info: Dict, stats: Dict, output_path: str, accumulated_content: str = None):
    """Export Word -"""
    
    try:
        doc = Document()
        
        # PAGE DE GARDE
        doc.add_paragraph('\n' * 3)
        
        title = doc.add_heading("RAPPORT D'ANALYSE EDA", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(36)
            run.font.color.rgb = RGBColor(25, 55, 109)
            run.font.bold = True
        
        subtitle = doc.add_paragraph('Analyse Exploratoire des Données')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(70, 130, 180)
            run.font.italic = True
        
        doc.add_paragraph('\n' * 2)
        
        file_info = doc.add_paragraph()
        file_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        file_run = file_info.add_run(f"📄 {data_info.get('filename', 'N/A')}\n\n")
        file_run.font.size = Pt(14)
        file_run.font.bold = True
        file_run.font.color.rgb = RGBColor(44, 62, 80)
        
        date_run = file_info.add_run(f"📅 {datetime.now().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(108, 117, 125)
        
        doc.add_paragraph('\n' * 5)
        
        separator = doc.add_paragraph('━' * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in separator.runs:
            run.font.color.rgb = RGBColor(41, 128, 185)
            run.font.size = Pt(12)
        
        doc.add_paragraph('\n' * 2)
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer1 = footer_para.add_run('Généré par ')
        footer1.font.size = Pt(11)
        footer1.font.color.rgb = RGBColor(149, 165, 166)
        footer1.font.italic = True
        
        footer2 = footer_para.add_run('EDA-Desk PRO')
        footer2.font.size = Pt(11)
        footer2.font.color.rgb = RGBColor(41, 128, 185)
        footer2.font.bold = True
        
        # PAGE 2 - SYNTHÈSE
        doc.add_page_break()
        
        section_title = doc.add_heading('SYNTHÈSE EXÉCUTIVE', 1)
        for run in section_title.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(25, 55, 109)
            run.font.bold = True
        
        doc.add_paragraph()
        
        # Tableau SANS fond de couleur (évite l'erreur)
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = 'MÉTRIQUE'
        hdr_cells[1].text = 'VALEUR'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        info_data = [
            ('📊 Dimensions', f"{data_info.get('rows', 0):,} × {data_info.get('columns', 0)}"),
            ('🔢 Variables numériques', str(data_info.get('numeric_vars', 0))),
            ('🏷️ Variables catégorielles', str(data_info.get('categorical_vars', 0))),
            ('✓ Variables booléennes', str(data_info.get('boolean_vars', 0))),
            ('⭐ Score de qualité', f"{data_info.get('quality_score', 0):.1f} / 100"),
            ('📋 Nombre d\'analyses', str(data_info.get('analyses_count', 0))),
            ('📅 Date', datetime.now().strftime('%d/%m/%Y'))
        ]
        
        for i, (label, value) in enumerate(info_data, 1):
            row_cells = info_table.rows[i].cells
            
            row_cells[0].text = label
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(44, 62, 80)
            
            row_cells[1].text = value
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(41, 128, 185)
                    run.font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Évaluation
        quality_score = data_info.get('quality_score', 0)
        
        if quality_score >= 90:
            emoji, grade, grade_color = "🟢", "EXCELLENT", RGBColor(39, 174, 96)
        elif quality_score >= 75:
            emoji, grade, grade_color = "🟡", "BON", RGBColor(241, 196, 15)
        elif quality_score >= 60:
            emoji, grade, grade_color = "🟠", "MOYEN", RGBColor(230, 126, 34)
        else:
            emoji, grade, grade_color = "🔴", "À AMÉLIORER", RGBColor(231, 76, 60)
        
        eval_para = doc.add_paragraph()
        eval_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        eval_run = eval_para.add_run(f"\n{emoji}  ÉVALUATION : {grade}  {emoji}\n")
        eval_run.font.size = Pt(18)
        eval_run.font.bold = True
        eval_run.font.color.rgb = grade_color
        
        score_run = eval_para.add_run(f"\nScore : {quality_score:.1f} / 100\n")
        score_run.font.size = Pt(14)
        score_run.font.color.rgb = RGBColor(52, 73, 94)
        
        # ANALYSES DÉTAILLÉES
        if accumulated_content and data_info.get('analyses_count', 0) > 0:
            doc.add_page_break()
            
            analyses_title = doc.add_heading('ANALYSES DÉTAILLÉES', 1)
            for run in analyses_title.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(25, 55, 109)
                run.font.bold = True
            
            doc.add_paragraph()
            
            lines = accumulated_content.split('\n')
            analysis_count = 0
            
            for line in lines:
                line_stripped = line.strip()
                
                if not line_stripped:
                    continue
                
                if 'ANALYSE #' in line:
                    analysis_count += 1
                    doc.add_paragraph()
                    
                    num_para = doc.add_paragraph()
                    num_run = num_para.add_run(f"● Analyse {analysis_count}")
                    num_run.font.size = Pt(11)
                    num_run.font.color.rgb = RGBColor(41, 128, 185)
                    num_run.font.bold = True
                    
                    clean_title = line.replace('═', '').replace('ANALYSE #', '').strip()
                    if ' (à ' in clean_title:
                        clean_title = clean_title.split(' (à ')[0].strip()
                    
                    title_heading = doc.add_heading('', 2)
                    title_run = title_heading.add_run(f"▸ {clean_title}")
                    title_run.font.size = Pt(16)
                    title_run.font.color.rgb = RGBColor(41, 128, 185)
                    title_run.font.bold = True
                    
                    line_para = doc.add_paragraph('━' * 60)
                    for run in line_para.runs:
                        run.font.color.rgb = RGBColor(174, 214, 241)
                        run.font.size = Pt(10)
                    
                    continue
                
                if line_stripped.startswith('─') or line_stripped.startswith('═'):
                    continue
                
                if line_stripped:
                    content_para = doc.add_paragraph(line_stripped)
                    content_para.paragraph_format.left_indent = Inches(0.3)
                    for run in content_para.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(52, 73, 94)
        
        else:
            doc.add_page_break()
            
            diag_title = doc.add_heading('DIAGNOSTIC DE QUALITÉ', 1)
            for run in diag_title.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(25, 55, 109)
            
            doc.add_paragraph()
            
            metrics = [
                ('❓ Valeurs manquantes', f"{data_info.get('missing_pct', 0):.2f}%"),
                ('🎯 Outliers', str(data_info.get('outliers_count', 0))),
                ('📌 Variables constantes', str(data_info.get('constant_vars', 0)))
            ]
            
            for metric, value in metrics:
                p = doc.add_paragraph()
                
                m_run = p.add_run(f"{metric} : ")
                m_run.font.bold = True
                m_run.font.size = Pt(11)
                m_run.font.color.rgb = RGBColor(44, 62, 80)
                
                v_run = p.add_run(value)
                v_run.font.size = Pt(11)
                v_run.font.color.rgb = RGBColor(41, 128, 185)
                v_run.font.bold = True
        
        # FOOTER
        doc.add_page_break()
        doc.add_paragraph('\n' * 10)
        
        final_line = doc.add_paragraph('━' * 55)
        final_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in final_line.runs:
            run.font.color.rgb = RGBColor(41, 128, 185)
            run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        footer_final = doc.add_paragraph()
        footer_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        f1 = footer_final.add_run('Rapport généré par ')
        f1.font.size = Pt(12)
        f1.font.color.rgb = RGBColor(127, 140, 141)
        f1.font.italic = True
        
        f2 = footer_final.add_run('EDA-Desk PRO')
        f2.font.size = Pt(12)
        f2.font.color.rgb = RGBColor(41, 128, 185)
        f2.font.bold = True
        
        f3 = footer_final.add_run(f"\n{datetime.now().strftime('%d/%m/%Y')}")
        f3.font.size = Pt(10)
        f3.font.color.rgb = RGBColor(149, 165, 166)
        f3.font.italic = True
        
        version_para = doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        v1 = version_para.add_run('\n\nVersion 4.0 Professional\n')
        v1.font.size = Pt(9)
        v1.font.color.rgb = RGBColor(189, 195, 199)
        
        v2 = version_para.add_run('© 2024 EDA-Desk')
        v2.font.size = Pt(8)
        v2.font.color.rgb = RGBColor(189, 195, 199)
        v2.font.italic = True
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Erreur export Word: {e}")
        import traceback
        traceback.print_exc()
        return False
# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT PDF - DESIGN PROFESSIONNEL BLEU MASTERCLASS
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_pdf_masterclass(data_info: Dict, stats: Dict, output_path: str, accumulated_content: str = None):
    """Export PDF avec design professionnel bleu élégant"""
    
    # Créer le document avec une classe personnalisée pour le header/footer
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            canvas.Canvas.__init__(self, *args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_number(num_pages)
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

        def draw_page_number(self, page_count):
            page = self._pageNumber
            if page > 1:  # Pas de numéro sur la page de garde
                # Ligne bleue en haut
                self.setStrokeColorRGB(0.16, 0.50, 0.73)  # Bleu corporate
                self.setLineWidth(2)
                self.line(50, 800, 545, 800)
                
                # Numéro de page
                self.setFont('Helvetica', 9)
                self.setFillColorRGB(0.5, 0.5, 0.5)
                self.drawRightString(545, 30, f"Page {page} / {page_count}")
                
                # Footer
                self.setFont('Helvetica-Oblique', 8)
                self.setFillColorRGB(0.7, 0.7, 0.7)
                self.drawString(50, 30, "EDA-Desk PRO")
    
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    # ═══════════════════════════════════════════════════════════════
    # STYLES PERSONNALISÉS BLEU ÉLÉGANT
    # ═══════════════════════════════════════════════════════════════
    
    # Style titre principal
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=36,
        textColor=colors.HexColor('#19376D'),  # Bleu marine
        spaceAfter=20,
        spaceBefore=100,
        alignment=1,  # Centre
        fontName='Helvetica-Bold',
        leading=42
    )
    
    # Style sous-titre
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=18,
        textColor=colors.HexColor('#4682B4'),  # Bleu acier
        spaceAfter=30,
        alignment=1,
        fontName='Helvetica-Oblique'
    )
    
    # Style info fichier
    fileinfo_style = ParagraphStyle(
        'FileInfo',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=10,
        alignment=1,
        fontName='Helvetica-Bold'
    )
    
    # Style heading section
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=18,
        textColor=colors.HexColor('#19376D'),
        spaceAfter=15,
        spaceBefore=20,
        fontName='Helvetica-Bold',
        borderWidth=0,
        borderPadding=8,
        leftIndent=0
    )
    
    # Style sous-heading
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=14,
        textColor=colors.HexColor('#2980B9'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    # Style texte normal
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#34495E'),
        fontName='Helvetica',
        leading=14
    )
    
    # Style monospace
    mono_style = ParagraphStyle(
        'MonoText',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=9,
        textColor=colors.HexColor('#2C3E50'),
        leftIndent=20,
        spaceAfter=5
    )
    
    # ═══════════════════════════════════════════════════════════════
    # PAGE DE GARDE ÉLÉGANTE
    # ═══════════════════════════════════════════════════════════════
    
    # Titre principal
    story.append(Paragraph("RAPPORT D'ANALYSE", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Sous-titre
    story.append(Paragraph("Analyse Exploratoire des Données", subtitle_style))
    story.append(Spacer(1, 0.5*inch))
    
    # Info fichier
    file_text = f"<b>📄 {data_info.get('filename', 'N/A')}</b>"
    story.append(Paragraph(file_text, fileinfo_style))
    
    # MODIFIÉ: Seulement la date
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#6C757D'),
        alignment=1
    )
    story.append(Paragraph(f"📅 {datetime.now().strftime('%d %B %Y')}", date_style))
    
    story.append(Spacer(1, 1.5*inch))
    
    # Ligne de séparation
    separator_line = Paragraph("━" * 70, ParagraphStyle(
        'Separator',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#2980B9'),
        alignment=1
    ))
    story.append(separator_line)
    
    story.append(Spacer(1, 1*inch))
    
    # Footer page de garde
    footer_style = ParagraphStyle(
        'FooterCover',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#95A5A6'),
        alignment=1,
        fontName='Helvetica-Oblique'
    )
    story.append(Paragraph("Généré par <b><font color='#2980B9'>EDA-Desk PRO</font></b>", footer_style))
    
    # ═══════════════════════════════════════════════════════════════
    # PAGE 2 - SYNTHÈSE EXÉCUTIVE
    # ═══════════════════════════════════════════════════════════════
    
    story.append(PageBreak())
    
    # Titre section
    story.append(Paragraph("■ SYNTHÈSE EXÉCUTIVE", heading_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Tableau des métriques avec design moderne
    data_table = [
        ['MÉTRIQUES CLÉS', 'VALEUR']
    ]
    
    # MODIFIÉ: Seulement la date
    info_data = [
        ('📊 Dimensions du dataset', f"{data_info.get('rows', 0):,} lignes × {data_info.get('columns', 0)} colonnes"),
        ('🔢 Variables numériques', str(data_info.get('numeric_vars', 0))),
        ('🏷️ Variables catégorielles', str(data_info.get('categorical_vars', 0))),
        ('✓ Variables booléennes', str(data_info.get('boolean_vars', 0))),
        ('⭐ Score de qualité', f"{data_info.get('quality_score', 0):.1f} / 100"),
        ('📋 Nombre d\'analyses', str(data_info.get('analyses_count', 0))),
        ('📅 Date de l\'analyse', datetime.now().strftime('%d/%m/%Y'))
    ]
    
    for label, value in info_data:
        data_table.append([label, value])
    
    # Créer le tableau avec style élégant
    info_table = Table(data_table, colWidths=[3*inch, 3*inch])
    
    table_style = TableStyle([
        # En-tête avec fond bleu marine
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#19376D')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
        ('TOPPADDING', (0, 0), (-1, 0), 15),
        
        # Labels (colonne 1)
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#EBF5FB')),  # Bleu glacier
        ('TEXTCOLOR', (0, 1), (0, -1), colors.HexColor('#2C3E50')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (0, -1), 10),
        ('LEFTPADDING', (0, 1), (0, -1), 15),
        
        # Valeurs (colonne 2)
        ('TEXTCOLOR', (1, 1), (1, -1), colors.HexColor('#2980B9')),
        ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 1), (1, -1), 10),
        ('RIGHTPADDING', (1, 1), (1, -1), 15),
        
        # Alternance de fond
        ('ROWBACKGROUNDS', (1, 1), (1, -1), [colors.white, colors.HexColor('#F8F9FA')]),
        
        # Bordures élégantes
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
        ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#2980B9')),
        
        # Padding
        ('TOPPADDING', (0, 1), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 12),
    ])
    
    info_table.setStyle(table_style)
    story.append(info_table)
    
    story.append(Spacer(1, 0.4*inch))
    
    # Card d'évaluation globale
    quality_score = data_info.get('quality_score', 0)
    
    if quality_score >= 90:
        emoji = "🟢"
        grade = "EXCELLENT"
        grade_color = '#27AE60'
        grade_bg = '#D5F4E6'
    elif quality_score >= 75:
        emoji = "🟡"
        grade = "BON"
        grade_color = '#F1C40F'
        grade_bg = '#FCF3CF'
    elif quality_score >= 60:
        emoji = "🟠"
        grade = "MOYEN"
        grade_color = '#E67E22'
        grade_bg = '#FAE5D3'
    else:
        emoji = "🔴"
        grade = "À AMÉLIORER"
        grade_color = '#E74C3C'
        grade_bg = '#FADBD8'
    
    # Tableau pour la card d'évaluation
    eval_data = [[f"{emoji}  ÉVALUATION GLOBALE : {grade}  {emoji}"]]
    eval_table = Table(eval_data, colWidths=[6*inch])
    
    eval_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(grade_bg)),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor(grade_color)),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 16),
        ('TOPPADDING', (0, 0), (-1, -1), 20),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ('BOX', (0, 0), (-1, -1), 2, colors.HexColor(grade_color)),
    ])
    
    eval_table.setStyle(eval_style)
    story.append(eval_table)
    
    # Score détaillé
    score_text = f"<b>Score de qualité : {quality_score:.1f} / 100</b>"
    score_para_style = ParagraphStyle(
        'ScoreDetail',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#34495E'),
        alignment=1,
        spaceBefore=10
    )
    story.append(Paragraph(score_text, score_para_style))
    
    # ═══════════════════════════════════════════════════════════════
    # ANALYSES DÉTAILLÉES
    # ═══════════════════════════════════════════════════════════════
    
    if accumulated_content and data_info.get('analyses_count', 0) > 0:
        story.append(PageBreak())
        story.append(Paragraph("■ ANALYSES DÉTAILLÉES", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Parser le contenu
        lines = accumulated_content.split('\n')
        analysis_count = 0
        current_section = []
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped or line_stripped.startswith('╔') or line_stripped.startswith('╚'):
                continue
            
            # Titre d'analyse - MODIFIÉ: Supprimer l'heure
            if 'ANALYSE #' in line:
                analysis_count += 1
                
                # Ajouter section précédente
                if current_section:
                    section_text = '<br/>'.join(current_section)
                    story.append(Paragraph(section_text, mono_style))
                    current_section = []
                
                story.append(Spacer(1, 0.15*inch))
                
                # Numéro
                num_style = ParagraphStyle(
                    'AnalysisNum',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=colors.HexColor('#2980B9'),
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(f"● Analyse {analysis_count}", num_style))
                
                # Titre - NETTOYER pour enlever l'heure
                clean_title = line.replace('═', '').replace('ANALYSE #', '').replace(str(analysis_count), '').strip()
                
                # Supprimer la partie avec l'heure "(à HH:MM:SS)"
                if ' (à ' in clean_title:
                    clean_title = clean_title.split(' (à ')[0].strip()
                
                story.append(Paragraph(f"▸ {clean_title}", subheading_style))
                
                # Ligne
                line_para = Paragraph("━" * 60, ParagraphStyle(
                    'LineDeco',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=colors.HexColor('#AED6F1')
                ))
                story.append(line_para)
                story.append(Spacer(1, 0.1*inch))
                continue
            
            # Contenu
            if line_stripped.startswith('║'):
                clean_line = line.replace('║', '').strip()
                if clean_line:
                    current_section.append(clean_line)
                continue
            
            if line_stripped.startswith('─') or line_stripped.startswith('═'):
                if current_section:
                    section_text = '<br/>'.join(current_section)
                    story.append(Paragraph(section_text, mono_style))
                    current_section = []
                story.append(Spacer(1, 0.05*inch))
                continue
            
            if line_stripped:
                current_section.append(line_stripped)
        
        # Dernière section
        if current_section:
            section_text = '<br/>'.join(current_section)
            story.append(Paragraph(section_text, mono_style))
    
    else:
        # Fallback - Diagnostic basique
        story.append(PageBreak())
        story.append(Paragraph("■ DIAGNOSTIC DE QUALITÉ", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        metrics_data = [
            ['MÉTRIQUE', 'VALEUR'],
            ['❓ Valeurs manquantes', f"{data_info.get('missing_pct', 0):.2f}%"],
            ['🎯 Outliers détectés', str(data_info.get('outliers_count', 0))],
            ['📌 Variables constantes', str(data_info.get('constant_vars', 0))]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 3*inch])
        
        metrics_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#19376D')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
            ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#2980B9')),
        ])
        
        metrics_table.setStyle(metrics_style)
        story.append(metrics_table)
    
    # ═══════════════════════════════════════════════════════════════
    # PAGE FINALE - Footer
    # ═══════════════════════════════════════════════════════════════
    
    story.append(PageBreak())
    story.append(Spacer(1, 4*inch))
    
    # Ligne finale
    final_line = Paragraph("━" * 55, ParagraphStyle(
        'FinalLine',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor('#2980B9'),
        alignment=1
    ))
    story.append(final_line)
    story.append(Spacer(1, 0.3*inch))
    
    # Footer - MODIFIÉ: Seulement la date
    footer_final_style = ParagraphStyle(
        'FooterFinal',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#7F8C8D'),
        alignment=1,
        fontName='Helvetica-Oblique'
    )
    
    footer_text = f"Rapport généré par <b><font color='#2980B9'>EDA-Desk PRO</font></b><br/>{datetime.now().strftime('%d %B %Y')}"
    story.append(Paragraph(footer_text, footer_final_style))
    
    story.append(Spacer(1, 0.3*inch))
    
    # Version
    version_style = ParagraphStyle(
        'Version',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#BDC3C7'),
        alignment=1
    )
    story.append(Paragraph("Version 4.0 Professional Edition", version_style))
    story.append(Paragraph("© 2024 EDA-Desk - Tous droits réservés", version_style))
    
    # Build PDF avec canvas personnalisé
    pdf.build(story, canvasmaker=NumberedCanvas)